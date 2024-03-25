import argparse
import re
from docx import Document

def fjern_tekst_mellom(doc, pattern):
    # Fjern fra paragrafer
    for paragraph in doc.paragraphs:
        if pattern.search(paragraph.text):
            paragraph.text = pattern.sub('', paragraph.text)

    # Fjern fra tabellceller
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if pattern.search(cell.text):
                    cell.text = pattern.sub('', cell.text)

def fjern_innhold(docx_path, ny_docx_path):
    doc = Document(docx_path)
    pattern = re.compile(r'med \w{1,4} konsulenter')

    fjern_tekst_mellom(doc, pattern)

    doc.save(ny_docx_path)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Fjerner tekst mellom "med" og "konsulenter" i en .docx-fil, inkludert innenfor tabeller.')
    parser.add_argument('input_file', type=str, help='Stien til input .docx-fil.')
    parser.add_argument('output_file', type=str, help='Stien til output .docx-fil.')

    args = parser.parse_args()

    fjern_innhold(args.input_file, args.output_file)

    print(f"Filbehandlingen er fullført. Sjekk '{args.output_file}' for resultatet.")
