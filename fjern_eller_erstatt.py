import argparse
import re
from docx import Document

def fjern_eller_erstatt_tekst(doc, soke_pattern, erstatt_med=None):
    compiled_pattern = re.compile(soke_pattern)

    # Fjern eller erstatt i paragrafer
    for paragraph in doc.paragraphs:
        if compiled_pattern.search(paragraph.text):
            paragraph.text = compiled_pattern.sub(erstatt_med if erstatt_med else '', paragraph.text)

    # Fjern eller erstatt i tabellceller
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if compiled_pattern.search(cell.text):
                    cell.text = compiled_pattern.sub(erstatt_med if erstatt_med else '', cell.text)

def bearbeid_docx(docx_path, ny_docx_path, soke_pattern, erstatt_med=None):
    doc = Document(docx_path)
    fjern_eller_erstatt_tekst(doc, soke_pattern, erstatt_med)
    doc.save(ny_docx_path)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Søker og erstatter (eller fjerner) tekst i en .docx-fil, inkludert innenfor tabeller.')
    parser.add_argument('input_file', type=str, help='Stien til input .docx-fil.')
    parser.add_argument('output_file', type=str, help='Stien til output .docx-fil.')
    parser.add_argument('soke_pattern', type=str, help='Regulært uttrykk for tekst som skal søkes etter.')
    parser.add_argument('--erstatt_med', type=str, default='', help='Teksten som skal erstatte det søkte mønsteret. Hvis ikke angitt, fjernes den søkte teksten.')

    args = parser.parse_args()

    bearbeid_docx(args.input_file, args.output_file, args.soke_pattern, args.erstatt_med)

    print(f"Filbehandlingen er fullført. Sjekk '{args.output_file}' for resultatet.")

